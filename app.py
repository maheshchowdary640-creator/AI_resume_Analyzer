import os
import re
import json
import logging
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import uvicorn

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("AIResumeAnalyzer")

app = FastAPI(title="AI Resume Analyzer & Job Matcher API")

# Add CORS Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Text extraction helpers
def extract_text_from_pdf(file_bytes: bytes) -> str:
    import io
    from pypdf import PdfReader
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    import io
    from docx import Document
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

# Mock data engines
def generate_mock_resume_data(raw_text: str) -> Dict[str, Any]:
    # Extract email and phone if possible
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', raw_text)
    phone_match = re.search(r'\(?\+?\d{1,4}\)?[-\s\.]?\d{3}[-\s\.]?\d{3,4}[-\s\.]?\d{3,4}', raw_text)
    
    email = email_match.group(0) if email_match else "contact@example.com"
    phone = phone_match.group(0) if phone_match else "+1 (555) 019-2834"
    
    # Extract candidate name (heuristics: first line if short and non-numeric)
    lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
    name = "Jane Doe"
    if lines:
        for line in lines[:3]:
            if len(line.split()) <= 4 and not any(char.isdigit() for char in line) and "@" not in line and "curriculum" not in line.lower() and "resume" not in line.lower():
                name = line
                break

    # Extract some skills from raw text to show partial parsing
    sample_skills = ["JavaScript", "Python", "React", "Node.js", "Docker", "Git", "SQL", "HTML/CSS"]
    found_skills = [s for s in sample_skills if s.lower() in raw_text.lower()]
    if not found_skills:
        found_skills = ["Python", "SQL", "Docker", "Git"]

    return {
        "name": name,
        "contact": {
            "email": email,
            "phone": phone,
            "linkedin": f"linkedin.com/in/{name.lower().replace(' ', '')}",
            "website": f"{name.lower().replace(' ', '')}.dev"
        },
        "professional_summary": "Highly motivated and results-driven Software Engineer with over 4 years of experience. Specializes in building highly scalable APIs, beautiful web interfaces, and modern cloud deployment flows. Passionate about solving complex architecture challenges and optimizing code performance.",
        "skills": {
            "technical": found_skills + ["FastAPI", "TypeScript", "PostgreSQL"],
            "soft": ["Problem Solving", "Collaboration", "Agile Methodologies", "Technical Writing"],
            "tools": ["AWS", "GitHub Actions", "VS Code", "Postman"]
        },
        "work_experience": [
            {
                "company": "Enterprise Software LLC",
                "role": "Software Engineer II",
                "start_date": "2023-05",
                "end_date": "Present",
                "description_bullets": [
                    "Developed and maintained high-performance Python APIs, improving overall system latency by 25%.",
                    "Redesigned the frontend dashboard using React, enhancing responsive design and user engagement metrics by 15%.",
                    "Configured unified testing pipelines, reducing deployment errors by 30%."
                ]
            },
            {
                "company": "Innovate Labs",
                "role": "Junior Developer",
                "start_date": "2021-06",
                "end_date": "2023-04",
                "description_bullets": [
                    "Built reusable UI components in JavaScript and CSS, adhering to accessibility standards.",
                    "Optimized SQL database indexes, reducing query runtimes by 200ms on average.",
                    "Collaborated in a 6-person scrum team to deliver product features weekly."
                ]
            }
        ],
        "education": [
            {
                "institution": "Tech University",
                "degree": "B.S. in Computer Science",
                "graduation_date": "2021-05",
                "gpa": "3.7"
            }
        ],
        "projects": [
            {
                "name": "Cloud Tasks Tracker",
                "technologies": ["FastAPI", "PostgreSQL", "React"],
                "description": "A serverless task-tracking dashboard supporting real-time webhooks and OAuth login."
            }
        ],
        "certifications": [
            {
                "name": "AWS Certified Developer",
                "issuing_organization": "Amazon Web Services",
                "date": "2024-01"
            }
        ]
    }

def generate_mock_match_data(resume_data: Dict[str, Any], job_description: str) -> Dict[str, Any]:
    job_lower = job_description.lower()
    
    resume_skills = resume_data.get("skills", {}).get("technical", [])
    if not resume_skills:
        resume_skills = ["Python", "SQL", "Git", "React"]
        
    matching_skills = []
    missing_skills = []
    
    # Mocking standard keywords to check
    keywords = ["React", "Python", "FastAPI", "Docker", "AWS", "SQL", "TypeScript", "Kubernetes", "GraphQL", "CI/CD"]
    
    for kw in keywords:
        kw_in_job = kw.lower() in job_lower
        kw_in_resume = any(kw.lower() in str(s).lower() for s in resume_skills)
        
        if kw_in_job:
            if kw_in_resume:
                matching_skills.append(kw)
            else:
                missing_skills.append(kw)
                
    # Safeguard if no keywords match job description
    if not matching_skills and not missing_skills:
        matching_skills = ["Python", "SQL"]
        missing_skills = ["AWS", "Docker", "FastAPI"]
        
    total_relevant = len(matching_skills) + len(missing_skills)
    score = int((len(matching_skills) / total_relevant) * 100) if total_relevant > 0 else 70
    score = max(40, min(95, score)) # Keep realistic boundaries
    
    keyword_analysis = []
    for kw in keywords:
        if kw.lower() in job_lower:
            keyword_analysis.append({
                "keyword": kw,
                "relevance_score": 9 if kw in ["React", "Python", "FastAPI"] else 6,
                "status": "found" if kw in matching_skills else "missing"
            })
            
    if not keyword_analysis:
        keyword_analysis = [
            {"keyword": "Software Engineering", "relevance_score": 9, "status": "found"},
            {"keyword": "APIs", "relevance_score": 8, "status": "found"},
            {"keyword": "Cloud", "relevance_score": 7, "status": "missing"},
        ]

    return {
        "match_score": score,
        "skills_analysis": {
            "matching_skills": matching_skills,
            "missing_skills": missing_skills
        },
        "keyword_analysis": keyword_analysis,
        "resume_feedback": [
            "Your resume features strong core technologies but lacks explicit mention of cloud containerization matching the job description.",
            "Incorporate metrics-oriented achievements in your enterprise work experience to prove business impact.",
            "Include your certification directly near the header or skills category to catch recruiters' eyes."
        ],
        "tailored_suggestions": {
            "summary": f"Experienced Developer skilled in {' / '.join(matching_skills[:3])}. Proven track record of designing REST APIs and leading agile implementations. Seeking to apply software expertise to achieve high-performance results for the role.",
            "experience": [
                f"Add: 'Leveraged {' and '.join(matching_skills[:2])} to build highly scalable background services, meeting job specifications.'",
                "Add: 'Containerized deployment configurations (Docker), resolving local dependency issues and aligning staging environments.'"
            ],
            "projects": [
                f"Under Cloud Tasks Tracker, explicitly detail how you optimized DB queries to handle scale using AWS."
            ]
        }
    }

# Gemini API Integration helpers
def get_gemini_model(api_key: str):
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    
    model_name = "gemini-1.5-flash"
    try:
        # Dynamically query available models for this specific API key
        available_models = [m.name for m in genai.list_models()]
        logger.info(f"Available models for key: {available_models}")
        
        # Order of preference (using standard prefixed names)
        preferences = [
            "models/gemini-1.5-flash",
            "models/gemini-1.5-flash-latest",
            "models/gemini-1.5-pro",
            "models/gemini-pro"
        ]
        for pref in preferences:
            if pref in available_models:
                model_name = pref.replace("models/", "")
                break
        else:
            # Fallback to the first model supporting generateContent
            for m in genai.list_models():
                if "generateContent" in m.supported_generation_methods:
                    model_name = m.name.replace("models/", "")
                    break
    except Exception as e:
        logger.error(f"Failed to dynamically list models: {str(e)}. Defaulting to gemini-1.5-flash.")
        
    logger.info(f"Initializing GenerativeModel with: {model_name}")
    return genai.GenerativeModel(model_name)

# API Endpoints
@app.post("/api/parse")
async def parse_resume(
    file: UploadFile = File(...),
    x_api_key: Optional[str] = Header(None)
):
    try:
        filename = file.filename
        file_bytes = await file.read()
        
        # Determine file type
        ext = os.path.splitext(filename)[1].lower()
        raw_text = ""
        
        if ext == ".pdf":
            raw_text = extract_text_from_pdf(file_bytes)
        elif ext == ".docx":
            raw_text = extract_text_from_docx(file_bytes)
        elif ext in [".txt", ".json"]:
            raw_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT.")
            
        if not raw_text.strip():
            raise HTTPException(status_code=400, detail="The uploaded file appears to be empty or unreadable.")
            
        # Determine API Key: Header takes priority over environment
        api_key = x_api_key or os.environ.get("GEMINI_API_KEY")
        
        if not api_key or api_key == "null" or api_key == "undefined":
            logger.info("No Gemini API key provided. Falling back to Mock Data Mode.")
            parsed_data = generate_mock_resume_data(raw_text)
            parsed_data["_mode"] = "mock"
            return JSONResponse(content=parsed_data)
            
        # Call Gemini API
        try:
            model = get_gemini_model(api_key)
            prompt = f"""
            You are an expert ATS (Applicant Tracking System) parser and resume engineer.
            Analyze the following raw resume text and extract all information into a structured JSON object.
            Do not wrap in markdown tags like ```json. Return ONLY the raw JSON string matching the structure.

            Resume text:
            {raw_text}

            JSON Schema to match:
            {{
              "name": "Full Name",
              "contact": {{
                "email": "Email address",
                "phone": "Phone number",
                "linkedin": "LinkedIn profile URL",
                "website": "Portfolio website URL"
              }},
              "professional_summary": "A concise summary of their professional background.",
              "skills": {{
                "technical": ["list of technical skills like languages, frameworks, databases"],
                "soft": ["list of soft skills"],
                "tools": ["list of tools/platforms like Git, Docker, AWS, Figma"]
              }},
              "work_experience": [
                {{
                  "company": "Company Name",
                  "role": "Job Title",
                  "start_date": "YYYY-MM or Mon YYYY",
                  "end_date": "YYYY-MM or Present",
                  "description_bullets": ["Detailed achievement bullet points"]
                }}
              ],
              "education": [
                {{
                  "institution": "University/School Name",
                  "degree": "Degree and Major",
                  "graduation_date": "Graduation date",
                  "gpa": "GPA (if available)"
                }}
              ],
              "projects": [
                {{
                  "name": "Project Name",
                  "technologies": ["tech used"],
                  "description": "Short description of what was built and achieved"
                }}
              ],
              "certifications": [
                {{
                  "name": "Certification Name",
                  "issuing_organization": "Issuer",
                  "date": "Issue Date"
                }}
              ]
            }}
            """
            
            response = model.generate_content(
                prompt,
                generation_config={"response_mime_type": "application/json"}
            )
            
            # Parse output
            cleaned_response = response.text.strip()
            # Clean markdown code blocks if Gemini output contains them
            if cleaned_response.startswith("```"):
                cleaned_response = re.sub(r"^```[a-zA-Z]*\n", "", cleaned_response)
                cleaned_response = re.sub(r"\n```$", "", cleaned_response)
                
            parsed_json = json.loads(cleaned_response)
            parsed_json["_mode"] = "live"
            return JSONResponse(content=parsed_json)
            
        except Exception as e:
            logger.error(f"Gemini API parse failed: {str(e)}. Falling back to Mock Data Mode.")
            parsed_data = generate_mock_resume_data(raw_text)
            parsed_data["_mode"] = "mock_fallback"
            parsed_data["_error"] = str(e)
            return JSONResponse(content=parsed_data)
            
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")

class MatchRequest(BaseModel):
    resume_data: Dict[str, Any]
    job_description: str

@app.post("/api/match")
async def match_resume_to_job(
    request: MatchRequest,
    x_api_key: Optional[str] = Header(None)
):
    try:
        api_key = x_api_key or os.environ.get("GEMINI_API_KEY")
        
        if not api_key or api_key == "null" or api_key == "undefined":
            logger.info("No Gemini API key provided. Falling back to Mock Matching Mode.")
            match_data = generate_mock_match_data(request.resume_data, request.job_description)
            match_data["_mode"] = "mock"
            return JSONResponse(content=match_data)
            
        try:
            model = get_gemini_model(api_key)
            prompt = f"""
            You are an expert technical recruiter and resume coach.
            Compare the following resume data against the job description.
            Analyze the alignment, calculate an overall match score (0-100), identify matching and missing skills, analyze keyword density relevance, provide detailed resume feedback, and suggest tailored improvements.

            Do not wrap in markdown code tags. Return ONLY the raw JSON string matching the structure.

            Resume Data:
            {json.dumps(request.resume_data, indent=2)}

            Job Description:
            {request.job_description}

            JSON Schema to match:
            {{
              "match_score": 85, // integer from 0 to 100
              "skills_analysis": {{
                "matching_skills": ["skills present in both"],
                "missing_skills": ["critical skills from the job description that are missing from the resume"]
              }},
              "keyword_analysis": [
                {{
                  "keyword": "React", 
                  "relevance_score": 9, // importance in job description, 1 to 10
                  "status": "found" // "found" or "missing"
                }}
              ],
              "resume_feedback": [
                "Feedback point 1 regarding structure or content",
                "Feedback point 2 regarding tone or formatting"
              ],
              "tailored_suggestions": {{
                "summary": "Suggested tailored professional summary rewrite",
                "experience": [
                  "Rewritten experience bullet point 1 emphasizing target keywords",
                  "Rewritten experience bullet point 2 emphasizing target keywords"
                ],
                "projects": [
                  "Suggestion for a project to add or how to reframe an existing project"
                ]
              }}
            }}
            """
            
            response = model.generate_content(
                prompt,
                generation_config={"response_mime_type": "application/json"}
            )
            
            cleaned_response = response.text.strip()
            if cleaned_response.startswith("```"):
                cleaned_response = re.sub(r"^```[a-zA-Z]*\n", "", cleaned_response)
                cleaned_response = re.sub(r"\n```$", "", cleaned_response)
                
            match_json = json.loads(cleaned_response)
            match_json["_mode"] = "live"
            return JSONResponse(content=match_json)
            
        except Exception as e:
            logger.error(f"Gemini API match failed: {str(e)}. Falling back to Mock Match Mode.")
            match_data = generate_mock_match_data(request.resume_data, request.job_description)
            match_data["_mode"] = "mock_fallback"
            match_data["_error"] = str(e)
            return JSONResponse(content=match_data)
            
    except Exception as e:
        logger.error(f"Match endpoint failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")

class ChatMessage(BaseModel):
    role: str # "user" or "assistant"
    content: str

class ChatRequest(BaseModel):
    messages: List[ChatMessage]
    resume_data: Optional[Dict[str, Any]] = None
    job_description: Optional[str] = None

@app.post("/api/chat")
async def chat_assistant(
    request: ChatRequest,
    x_api_key: Optional[str] = Header(None)
):
    try:
        api_key = x_api_key or os.environ.get("GEMINI_API_KEY")
        
        user_message = request.messages[-1].content
        
        if not api_key or api_key == "null" or api_key == "undefined":
            logger.info("No Gemini API key provided. Falling back to Mock Chat Mode.")
            # Simple mock chatbot replies
            reply = ""
            user_msg_lower = user_message.lower()
            
            skills_list = request.resume_data.get("skills", {}).get("technical", []) if request.resume_data else []
            
            if "skill" in user_msg_lower:
                if skills_list:
                    reply = f"Looking at your profile, your technical skills are: **{', '.join(skills_list)}**. "
                    if request.job_description:
                        reply += "To better match the job, you might want to add some missing keywords such as **AWS** or **Docker** if you have experience in them. "
                else:
                    reply = "I don't see a resume loaded yet. Please upload your resume in the **Resume Analyzer** tab so I can inspect your skills!"
            elif "experience" in user_msg_lower or "work" in user_msg_lower:
                reply = "Your work experience details are structured nicely! To make it look more impactful to recruiters, try framing your accomplishments using the **X-Y-Z formula**:\n\n* *'Accomplished [X], as measured by [Y], by doing [Z]'.*\n\nFor example, instead of 'Wrote API backends', write: *'Developed robust backend endpoints (Z) in Python, decreasing API latency by 25% (Y) as measured by server metrics (X)'.*"
            elif "project" in user_msg_lower:
                reply = "Projects are a great way to bridge the gap if you lack commercial experience in a specific framework. If you are targeting a modern backend stack, I'd suggest building a **Task Scheduler with FastAPI and Redis** to showcase your understanding of asynchronous queues."
            else:
                reply = "Hello! I am your AI career assistant. I have reviewed your current dashboard state. Feel free to ask me questions like:\n\n1. *\"How can I optimize my skills for this job?\"*\n2. *\"Can you critique my work experience section?\"*\n3. *\"What projects should I add to my resume?\"*"
                
            return JSONResponse(content={"reply": reply, "_mode": "mock"})
            
        try:
            model = get_gemini_model(api_key)
            
            # Format chat history for the prompt context
            chat_context = []
            # Feed the last 6 messages to keep context window clean
            for msg in request.messages[:-1]:
                role_label = "User" if msg.role == "user" else "Assistant"
                chat_context.append(f"{role_label}: {msg.content}")
                
            history_text = "\n".join(chat_context)
            
            # Setup context files
            resume_str = json.dumps(request.resume_data, indent=2) if request.resume_data else "No resume uploaded yet."
            job_str = request.job_description if request.job_description else "No job description linked yet."
            
            prompt = f"""
            You are a highly supportive, knowledgeable, and professional AI Career Assistant and Resume Coach.
            Your goal is to help the user optimize their resume, understand job matching results, prepare for interviews, and advance their career.

            You have access to the user's parsed resume and the target job description (if provided).
            
            [Parsed Resume Context]
            {resume_str}

            [Target Job Description Context]
            {job_str}

            [Conversation History]
            {history_text}

            User's Prompt: {user_message}

            Instructions:
            1. Always be constructive, encouraging, and highly specific.
            2. If they ask to optimize a bullet point, rewrite it using the X-Y-Z formula (Accomplished [X] as measured by [Y], by doing [Z]).
            3. Reference skills from the job description when recommending changes.
            4. Keep your formatting clean, using Markdown (bolding, lists, code block structure) for readability. Do not start your response with 'Assistant:' or similar labels, just reply directly.
            """
            
            response = model.generate_content(prompt)
            return JSONResponse(content={"reply": response.text.strip(), "_mode": "live"})
            
        except Exception as e:
            logger.error(f"Gemini API chat failed: {str(e)}. Falling back to Mock Chat Mode.")
            return JSONResponse(content={"reply": f"Sorry, I encountered an API error calling Gemini ({str(e)}). Let me assist you in Mock mode: How would you like me to help you structure your resume bullet points?", "_mode": "mock_fallback"})
            
    except Exception as e:
        logger.error(f"Chat endpoint failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")

# Mount Static Files & Serve SPA
static_path = os.path.join(os.path.dirname(__file__), "static")
if not os.path.exists(static_path):
    os.makedirs(static_path)

@app.get("/")
async def read_index():
    index_file = os.path.join(static_path, "index.html")
    if os.path.exists(index_file):
        response = FileResponse(index_file)
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    return JSONResponse(content={"message": "Frontend static files are missing! Please build static/index.html first."}, status_code=404)

# Mount static files (HTML, CSS, JS)
app.mount("/", StaticFiles(directory=static_path), name="static")

if __name__ == "__main__":
    uvicorn.run("app:app", host="127.0.0.1", port=8000, reload=True)
